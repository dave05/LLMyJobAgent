import os
import logging
from typing import Dict, List
import re
from docx import Document
from PyPDF2 import PdfReader
from pdfminer.high_level import extract_text
import spacy

logger = logging.getLogger(__name__)

class ResumeParser:
    def __init__(self):
        """Initialize the resume parser with NLP model"""
        self.nlp = spacy.load('en_core_web_lg')
        self.supported_formats = ['.docx', '.pdf']
        
    def parse_resume(self, file_path: str) -> Dict[str, str]:
        """Parse resume file and extract text"""
        try:
            file_ext = os.path.splitext(file_path)[1].lower()
            
            if file_ext not in self.supported_formats:
                raise ValueError(f"Unsupported file format: {file_ext}")
                
            if file_ext == '.docx':
                return self._parse_docx(file_path)
            else:  # .pdf
                return self._parse_pdf(file_path)
                
        except Exception as e:
            logger.error(f"Error parsing resume: {str(e)}")
            raise
            
    def _parse_docx(self, file_path: str) -> Dict[str, str]:
        """Extract text from DOCX file"""
        try:
            doc = Document(file_path)
            text = []
            
            # Extract text from paragraphs
            for para in doc.paragraphs:
                text.append(para.text)
                
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text.append(cell.text)
                        
            return self._process_text('\n'.join(text))
            
        except Exception as e:
            logger.error(f"Error parsing DOCX: {str(e)}")
            raise
            
    def _parse_pdf(self, file_path: str) -> Dict[str, str]:
        """Extract text from PDF file"""
        try:
            # First try with PyPDF2
            try:
                reader = PdfReader(file_path)
                text = []
                for page in reader.pages:
                    text.append(page.extract_text())
                pdf_text = '\n'.join(text)
            except Exception:
                # Fallback to pdfminer if PyPDF2 fails
                pdf_text = extract_text(file_path)
                
            return self._process_text(pdf_text)
            
        except Exception as e:
            logger.error(f"Error parsing PDF: {str(e)}")
            raise
            
    def _process_text(self, text: str) -> Dict[str, str]:
        """Process extracted text and identify sections"""
        # Clean text
        text = self._clean_text(text)
        
        # Extract sections
        sections = self._extract_sections(text)
        
        # Process each section with NLP
        processed_sections = {}
        for section_name, section_text in sections.items():
            doc = self.nlp(section_text)
            processed_sections[section_name] = {
                'text': section_text,
                'entities': [(ent.text, ent.label_) for ent in doc.ents],
                'keywords': [token.text for token in doc if token.is_alpha and not token.is_stop]
            }
            
        return processed_sections
        
    def _clean_text(self, text: str) -> str:
        """Clean and normalize text"""
        # Remove extra whitespace
        text = re.sub(r'\s+', ' ', text)
        
        # Remove special characters but keep punctuation
        text = re.sub(r'[^\w\s.,;:!?()\-]', ' ', text)
        
        # Normalize line breaks
        text = re.sub(r'\n+', '\n', text)
        
        return text.strip()
        
    def _extract_sections(self, text: str) -> Dict[str, str]:
        """Extract common resume sections"""
        sections = {
            'contact': '',
            'summary': '',
            'experience': '',
            'education': '',
            'skills': '',
            'other': ''
        }
        
        # Common section headers
        section_patterns = {
            'contact': r'(?i)(?:contact|personal)\s*(?:information|details)?',
            'summary': r'(?i)(?:summary|profile|objective)',
            'experience': r'(?i)(?:experience|work\s*history|employment)',
            'education': r'(?i)(?:education|academic)',
            'skills': r'(?i)(?:skills|technical\s*skills|competencies)'
        }
        
        # Split text into lines
        lines = text.split('\n')
        current_section = 'other'
        
        for line in lines:
            # Check if line matches any section header
            for section, pattern in section_patterns.items():
                if re.search(pattern, line):
                    current_section = section
                    break
                    
            # Add line to current section
            if current_section in sections:
                sections[current_section] += line + '\n'
                
        # Clean each section
        for section in sections:
            sections[section] = self._clean_text(sections[section])
            
        return sections